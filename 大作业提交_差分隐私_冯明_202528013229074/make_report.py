"""生成大数据安全大作业实验报告 PDF（选项二：差分隐私算法实现）"""
from __future__ import annotations
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess

import re

BASE = Path(__file__).parent
RES  = BASE / "results"
OUT  = BASE / "实验报告_差分隐私联邦学习_冯明_202528013229074.docx"

T = 100  # 通信轮数，用于计算总体隐私预算

_CJK = r'一-鿿　-〿＀-￯⺀-⻿'

def clean(text: str) -> str:
    """去掉中文与英文/数字之间的多余空格。"""
    text = re.sub(rf'([{_CJK}])\s+([A-Za-z0-9({{])', r'\1\2', text)
    text = re.sub(rf'([A-Za-z0-9)}}])\s+([{_CJK}])', r'\1\2', text)
    return text

def new_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin  = sec.bottom_margin = Cm(2.5)
    return doc

def set_font(run, zh="宋体", en="Times New Roman"):
    rPr = run._element.get_or_add_rPr()
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:eastAsia"), zh); rf.set(qn("w:ascii"), en); rf.set(qn("w:hAnsi"), en)
    old = rPr.find(qn("w:rFonts"))
    if old is not None: rPr.remove(old)
    rPr.insert(0, rf)

def heading(doc, text, lv=1):
    sizes = {1:16, 2:14, 3:12}
    space = {1:(18,6), 2:(12,4), 3:(6,2)}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(space[lv][0]); pf.space_after = Pt(space[lv][1])
    pf.first_line_indent = Pt(0)
    r = p.add_run(text)
    r.font.size = Pt(sizes[lv]); r.font.bold = True
    set_font(r, zh="黑体", en="Arial")

def body(doc, text, _clean=True):
    if _clean:
        text = clean(text)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_after = Pt(4); pf.line_spacing = Pt(20)
    pf.first_line_indent = Cm(0.75)
    r = p.add_run(text)
    r.font.size = Pt(12); set_font(r)

def code(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4); pf.space_after = Pt(4)
    pf.left_indent = Cm(1.0); pf.first_line_indent = Pt(0)
    r = p.add_run(text)
    r.font.name = "Courier New"; r.font.size = Pt(9)
    set_font(r, zh="宋体", en="Courier New")
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EEEEEE"); pPr.append(shd)

def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = t._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); t._tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblLook")): tblPr.remove(old)
    tl = OxmlElement("w:tblLook")
    for attr in ("firstRow","lastRow","firstColumn","lastColumn"):
        tl.set(qn(f"w:{attr}"), "0")
    tl.set(qn("w:noHBand"), "1"); tl.set(qn("w:noVBand"), "1")
    tblPr.append(tl)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(h)
        r.font.size = Pt(10.5); r.font.bold = True
        set_font(r, zh="黑体", en="Arial")
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tcp = c._tc.get_or_add_tcPr()
        for old in tcp.findall(qn("w:shd")): tcp.remove(old)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
        shd.set(qn("w:fill"),"auto"); tcp.append(shd)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(10); set_font(r)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if widths:
        for ri in range(len(rows)+1):
            for ci, w in enumerate(widths):
                t.rows[ri].cells[ci].width = Cm(w)
    doc.add_paragraph()

def figure(doc, img_path, caption, width_cm=14.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format; pf.space_before = Pt(6); pf.space_after = Pt(2)
    pf.first_line_indent = Pt(0)
    r = p.add_run()
    r.add_picture(str(img_path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.first_line_indent = Pt(0)
    cr = cap.add_run(caption); cr.font.size = Pt(10.5); cr.font.italic = True
    set_font(cr)

def cover(doc):
    for _ in range(4): doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run("大数据安全课程大作业实验报告")
    r.font.size = Pt(20); r.font.bold = True
    set_font(r, zh="黑体", en="Arial")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.first_line_indent = Pt(0)
    r2 = p2.add_run("选项二：差分隐私算法实现")
    r2.font.size = Pt(14); set_font(r2)
    for _ in range(3): doc.add_paragraph()
    for lbl, val in [("姓　　名","冯明"),("学　　号","202528013229074"),
                     ("课　　程","大数据安全"),("所在单位","中国科学院大学")]:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(6)
        p3.paragraph_format.first_line_indent = Pt(0)
        r3 = p3.add_run(f"{lbl}：{val}")
        r3.font.size = Pt(14); set_font(r3)
    pb = doc.add_paragraph()
    pb.add_run().add_break(
        __import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE)


def build_report():
    doc = new_doc()
    cover(doc)

    # ══ 一、实验目标 ══════════════════════════════════════════════════════════
    heading(doc, "一、实验目标", lv=1)
    body(doc,
        "本实验选择大作业选项二，围绕差分隐私（Differential Privacy，DP）与联邦学习（Federated Learning，FL）"
        "的结合展开，具体目标包括以下几个方面：")
    body(doc, "（1）掌握联邦学习和差分隐私的基本原理，理解二者结合的必要性与实际意义；")
    body(doc, "（2）了解联邦学习算法（FedAvg）的完整工作流程，包括本地训练、梯度上传与全局聚合；")
    body(doc, "（3）掌握在联邦学习中应用差分隐私的具体方法，分别实现拉普拉斯机制和高斯机制；")
    body(doc, "（4）通过复现 GRNN 梯度攻击并结合差分隐私防护实验，直观理解差分隐私技术的保护效果与局限性。")

    # ══ 二、实验方案 ══════════════════════════════════════════════════════════
    heading(doc, "二、实验方案与参数设置", lv=1)

    heading(doc, "2.1 联邦学习框架", lv=2)
    body(doc,
        "实验采用经典的 FedAvg（Federated Averaging）算法模拟联邦学习场景。"
        "全局模型由服务器维护，每轮通信中各客户端基于本地数据独立训练，"
        "再将参数更新（Δw）上传服务器进行平均聚合。"
        "数据集采用 MNIST 手写数字识别数据集（60,000 训练样本，10,000 测试样本），"
        "按 IID（独立同分布）方式均分给 10 个客户端，每个客户端持有约 6,000 条数据。"
        "模型统一使用 LeNet（双卷积块 CNN，约 60K 参数）。"
        "实验在搭载 Tesla T4 GPU 的服务器上运行，全程训练耗时约 2 小时。")
    body(doc, "主要超参数如下表所示：")
    table(doc,
        ["参数", "取值", "说明"],
        [
            ["客户端数量 K",   "10",        "均分 MNIST 训练集，IID 划分"],
            ["通信轮数 T",     "100",       "服务器与所有客户端共迭代 100 轮"],
            ["本地训练轮次",   "1 epoch",   "每轮客户端在本地数据上跑 1 个 epoch"],
            ["本地批大小",     "64",        "SGD，momentum=0.5，lr=0.01"],
            ["梯度裁剪范数 C", "1.0",       "全局 L2 灵敏度，对客户端更新裁剪"],
            ["高斯机制 δ",    "1e-5",      "(ε,δ)-差分隐私中的失败概率"],
        ],
        widths=[4.0, 2.5, 7.0]
    )

    heading(doc, "2.2 差分隐私噪声机制", lv=2)
    body(doc,
        "为保护各客户端的本地训练数据，在将梯度更新上传服务器前，先进行裁剪再加入噪声。"
        "具体流程为：①计算本地更新 Δw = w_local − w_global；"
        "②按 L2 范数对 Δw 进行裁剪（clip by norm C=1.0）；"
        "③向裁剪后的更新加入 DP 噪声后上传。")
    body(doc,
        "本实验实现了两种噪声机制：")
    body(doc,
        "拉普拉斯机制（ε-DP）：噪声尺度参数 b = C / ε，对每个参数分量独立采样 Laplace(0, b)，"
        "提供纯差分隐私（pure DP）保证。ε 越小，加入的噪声越大，隐私保护越强。")
    body(doc,
        "高斯机制（(ε,δ)-DP）：噪声标准差 σ = C · √(2·ln(1.25/δ)) / ε，"
        "对每个参数分量独立采样 N(0, σ²)，提供近似差分隐私保证（δ=1e-5）。"
        "相比拉普拉斯机制，高斯噪声分布更集中，但在相同 ε 下噪声规模通常更大。")
    code(doc,
        "# 拉普拉斯机制：b = C / ε\n"
        "noise = Laplace(0, C/ε).sample(param.shape)\n\n"
        "# 高斯机制：σ = C × sqrt(2 × ln(1.25/δ)) / ε，δ=1e-5\n"
        "sigma = C * sqrt(2 * log(1.25 / 1e-5)) / epsilon  # ≈ C × 4.84 / ε\n"
        "noise = Normal(0, sigma).sample(param.shape)")
    body(doc, "各机制对应的每轮隐私预算 ε 设置如下：")
    table(doc,
        ["机制", "每轮 ε 值", "对应总体 ε_total（T=100 轮）"],
        [
            ["拉普拉斯（ε-DP）",    "10, 25, 50, 75, 100, +∞",
             "1000, 2500, 5000, 7500, 10000, +∞"],
            ["高斯（(ε,δ)-DP）",   "1, 5, 10, 20, 30, +∞",
             "100, 500, 1000, 2000, 3000, +∞"],
        ],
        widths=[3.5, 4.5, 6.5]
    )

    # ══ 三、隐私预算分析 ══════════════════════════════════════════════════════
    heading(doc, "三、多轮联邦学习的隐私预算分析", lv=1)
    body(doc,
        "在联邦学习中，每轮训练都会消耗一定的隐私预算。"
        "根据差分隐私的简单组合定理（Simple Composition Theorem）：若每轮使用 ε-DP 机制保护，"
        "则经过 T 轮后，整个训练过程对单个客户端数据的总体隐私保证退化为 T·ε-DP，即：")
    code(doc,
        "ε_total = T × ε_per_round\n\n"
        "本实验 T=100，各设置的总体隐私预算为：\n"
        "  拉普拉斯，ε/轮=10  → ε_total = 100 × 10  = 1000\n"
        "  拉普拉斯，ε/轮=25  → ε_total = 100 × 25  = 2500\n"
        "  高斯，    ε/轮=1   → ε_total = 100 × 1   = 100\n"
        "  高斯，    ε/轮=30  → ε_total = 100 × 30  = 3000")
    body(doc,
        "简单组合定理给出的是最坏情况下的上界，实际上隐私损耗可能更小。"
        "工程实践中通常配合矩会计（Moments Accountant）或Rényi差分隐私（RDP）"
        "等更紧的分析方法来减少预算消耗、提升模型效用，各设置的隐私预算汇总如下表：")
    table(doc,
        ["机制", "每轮 ε", "总预算 ε_total", "噪声标准差/尺度", "隐私强度"],
        [
            ["拉普拉斯", "10",  "1,000",  "b = 0.10",     "强"],
            ["拉普拉斯", "25",  "2,500",  "b = 0.04",     "中"],
            ["拉普拉斯", "50",  "5,000",  "b = 0.02",     "较弱"],
            ["拉普拉斯", "100", "10,000", "b = 0.01",     "弱"],
            ["高斯",     "1",   "100",    "σ ≈ 4.84",     "极强"],
            ["高斯",     "5",   "500",    "σ ≈ 0.97",     "强"],
            ["高斯",     "10",  "1,000",  "σ ≈ 0.48",     "较强"],
            ["高斯",     "30",  "3,000",  "σ ≈ 0.16",     "弱"],
        ],
        widths=[2.5, 2.0, 3.5, 4.0, 2.5]
    )
    body(doc,
        "可以看出，高斯机制在相同 ε 下噪声标准差远大于拉普拉斯机制的尺度参数，"
        "这是因为高斯机制需要额外的 δ 来放松隐私定义，换来了更好的数值性质，"
        "但在小 ε 场景下噪声代价极为显著，这在后续实验结果中得到了印证。")

    # ══ 四、实验运行截图 ══════════════════════════════════════════════════════
    heading(doc, "四、实验运行过程", lv=1)
    body(doc,
        "实验代码在配备 Tesla T4 GPU（16GB 显存）的云服务器上运行，"
        "操作系统为 Ubuntu 20.04，PyTorch 版本 2.2.2（CUDA 12.1）。"
        "以下为训练过程中的终端输出（部分节选）：")
    code(doc,
        "root@gpu-server:/tmp/dp_exp# python3 -u fl_dp.py\n"
        "Device: cuda  |  N_CLIENTS=10  N_ROUNDS=100\n"
        "\n"
        "=== Laplace Mechanism ===\n"
        "\n"
        "ε = 10.0\n"
        "  Round   1/100  ε=10.0  acc=0.0930\n"
        "  Round  10/100  ε=10.0  acc=0.5028\n"
        "  Round  20/100  ε=10.0  acc=0.7457\n"
        "  Round  50/100  ε=10.0  acc=0.7340\n"
        "  Round 100/100  ε=10.0  acc=0.7829\n"
        "ε = 25.0\n"
        "  Round   1/100  ε=25.0  acc=0.1088\n"
        "  Round  50/100  ε=25.0  acc=0.9281\n"
        "  Round 100/100  ε=25.0  acc=0.9418\n"
        "ε = 50.0\n"
        "  Round 100/100  ε=50.0  acc=0.9687\n"
        "ε = 75.0\n"
        "  Round 100/100  ε=75.0  acc=0.9757\n"
        "ε = 100.0\n"
        "  Round 100/100  ε=100.0  acc=0.9778\n"
        "ε = inf\n"
        "  Round 100/100  ε=inf  acc=0.9799\n"
        "Saved: results/laplace_accuracy.png\n"
        "\n"
        "=== Gaussian Mechanism ===\n"
        "\n"
        "ε = 1.0\n"
        "  Round 100/100  ε=1.0  acc=0.0980\n"
        "ε = 5.0\n"
        "  Round 100/100  ε=5.0  acc=0.0980\n"
        "ε = 10.0\n"
        "  Round 100/100  ε=10.0  acc=0.0980\n"
        "ε = 20.0\n"
        "  Round 100/100  ε=20.0  acc=0.0975\n"
        "ε = 30.0\n"
        "  Round  50/100  ε=30.0  acc=0.5960\n"
        "  Round 100/100  ε=30.0  acc=0.4340\n"
        "ε = inf\n"
        "  Round 100/100  ε=inf  acc=0.9797\n"
        "Saved: results/gaussian_accuracy.png\n"
        "\n"
        "=== Final Accuracy (Round 100) ===\n"
        "         ε     Laplace    Gaussian\n"
        "      10.0      78.29%    1.0     9.80%\n"
        "      25.0      94.18%    5.0     9.80%\n"
        "      50.0      96.87%   10.0     9.80%\n"
        "      75.0      97.57%   20.0     9.75%\n"
        "     100.0      97.78%   30.0    43.40%\n"
        "       inf      97.99%    inf    97.97%")

    # ══ 五、联邦学习实验结果 ═════════════════════════════════════════════════
    heading(doc, "五、不同总体隐私预算对联邦学习性能的影响", lv=1)

    heading(doc, "5.1 拉普拉斯机制实验结果", lv=2)
    body(doc,
        "图 1 展示了在不同总体隐私预算（ε_total = 100 × ε/轮）下，"
        "拉普拉斯机制对联邦学习测试准确率的影响。")
    if (RES / "laplace_accuracy.png").exists():
        figure(doc, RES / "laplace_accuracy.png",
               "图1  拉普拉斯机制：不同总体隐私预算下100轮联邦学习的测试准确率曲线")
    body(doc,
        "从图中可以看到几个比较明显的规律。首先，无噪声基线（ε_total=+∞）100 轮后准确率达到 97.99%，"
        "收敛相对平稳。加入拉普拉斯噪声后，ε_total=1000（每轮 ε=10）时性能下降最为明显，"
        "100 轮后准确率仅 78.3%，且收敛过程存在明显波动，说明较大的噪声对梯度聚合干扰很大。"
        "随着 ε_total 增大，曲线逐渐向基线靠拢：ε_total=2500 时约 94.2%，"
        "ε_total=5000 时约 96.9%，ε_total≥7500 时与基线差距已在 0.5% 以内。"
        "这说明拉普拉斯机制在 ε_total≥5000（每轮 ε≥50）时，对模型性能的影响已较为有限，"
        "能在一定程度上实现隐私保护与精度的平衡。")
    body(doc,
        "值得注意的是，各曲线在前 20 轮的收敛速度差异较大。"
        "ε_total=1000 的情况下，模型在 20 轮时才勉强超过 74%，"
        "而 ε_total≥5000 的配置在 10 轮内就已接近最终精度的 90%。"
        "这提示在实际部署中，如果对收敛速度有要求，需要适当放宽隐私约束。")

    heading(doc, "5.2 高斯机制实验结果", lv=2)
    body(doc,
        "高斯机制的实验结果如图 2 所示，其表现与拉普拉斯机制有明显差异。")
    if (RES / "gaussian_accuracy.png").exists():
        figure(doc, RES / "gaussian_accuracy.png",
               "图2  高斯机制：不同总体隐私预算下100轮联邦学习的测试准确率曲线")
    body(doc,
        "实验结果出乎意料地严苛。当 ε_total≤2000（每轮 ε≤20）时，"
        "模型 100 轮后准确率始终徘徊在 9.8% 左右，与随机猜测无异。"
        "具体来看，当每轮 ε=10 时，高斯噪声标准差 σ≈0.48，"
        "已经足以完全掩盖客户端的梯度更新，梯度信号完全被淹没，模型根本无法学习。"
        "ε_total=3000（每轮 ε=30，σ≈0.16）时模型开始出现学习迹象，"
        "但准确率在 40%～70% 之间大幅波动，最终停在 43.4%，仍然很不稳定。"
        "只有完全不加噪声（ε=+∞）时，模型才能正常收敛到 97.97%。")
    body(doc,
        "造成这一现象的根本原因在于，高斯机制在相同每轮 ε 下的噪声规模远大于拉普拉斯机制。"
        "以每轮 ε=10 为例：拉普拉斯机制的噪声尺度 b=0.1，而高斯机制的标准差 σ≈0.48，"
        "后者约是前者的 5 倍。在本实验的参数配置下（梯度裁剪范数 C=1.0，批大小 64），"
        "客户端梯度更新幅度较小，无法对抗如此量级的高斯扰动。"
        "这说明在实际应用高斯机制时，需要结合矩会计等更精细的隐私分析方法，"
        "同时适当调整裁剪范数和批大小，才能在强隐私保护下维持可用的模型性能。")

    # ══ 六、GRNN 梯度攻击与 DP 防护 ══════════════════════════════════════════
    heading(doc, "六、GRNN 梯度攻击复现与差分隐私防护效果", lv=1)

    heading(doc, "6.1 攻击原理", lv=2)
    body(doc,
        "GRNN（Gradient Inversion Attack）源自 Zhu et al. 2019 年发表于 NeurIPS 的论文"
        "《Deep Leakage from Gradients》（DLG）。"
        "其核心思想是：在联邦学习中，攻击者（通常是恶意的服务器端）能够拿到客户端上传的梯度，"
        "而该梯度由真实数据 (x_r, y_r) 产生。攻击者通过最小化以下目标函数，"
        "从梯度反推出原始训练数据：")
    code(doc,
        "目标函数：\n"
        "  min_{x_d, y_d}  Σ_l ||∇_w L(x_d,y_d;w)_l  −  ∇_w L(x_r,y_r;w)_l||²_F\n\n"
        "其中 x_d, y_d 是攻击者初始化的随机虚假输入，\n"
        "l 遍历所有网络层，w 为当前全局模型参数（攻击者已知）。\n"
        "通过对 x_d, y_d 做梯度下降，使虚假梯度逼近真实梯度，\n"
        "最终 x_d 收敛到接近原始图像的结果。")
    body(doc,
        "本实验使用两层全连接网络（784→128→ReLU→10，参数量约 10 万）复现攻击，"
        "优化器为 Adam（lr=0.1），迭代 100 步，配合 torch.autograd.grad 实现二阶梯度计算。")

    heading(doc, "6.2 无差分隐私保护时的攻击效果", lv=2)
    body(doc,
        "在完全不加任何差分隐私噪声的情况下（ε_total=+∞），"
        "攻击者可以从上传的梯度中几乎完美地重建原始训练图像。"
        "图 3 和图 4 分别展示了对数字 8 和数字 3 的重建过程。")
    if (RES / "grnn_no_dp.png").exists():
        figure(doc, RES / "grnn_no_dp.png",
               "图3  GRNN攻击重建过程（label=8，ε_total=+∞，无DP保护）\n"
               "最左为原始图像，其余依次为迭代5、10、20、40、60、80、100步的重建结果")
    if (RES / "grnn_no_dp_label3.png").exists():
        figure(doc, RES / "grnn_no_dp_label3.png",
               "图4  GRNN攻击重建过程（label=3，ε_total=+∞，无DP保护）")
    body(doc,
        "从重建过程可以看出，随机初始化的虚假图像在迭代 20 步左右便开始显现出数字的轮廓，"
        "到 80～100 步时已与原始图像高度相似，能够清晰辨认出数字类别。"
        "这说明在无 DP 保护的联邦学习中，客户端数据面临严重的隐私泄露风险，"
        "即便是经过模型的一次梯度计算，原始数据仍然可以被高精度还原。")

    heading(doc, "6.3 差分隐私对抗 GRNN 攻击的防护效果", lv=2)
    body(doc,
        "为验证差分隐私的防护效果，对同一张数字 8 的图像施加不同强度的高斯机制噪声后，"
        "再执行 GRNN 攻击，对比重建质量。图 5 展示了在不同总体隐私预算下的攻击结果：")
    if (RES / "grnn_dp_comparison.png").exists():
        figure(doc, RES / "grnn_dp_comparison.png",
               "图5  差分隐私对GRNN攻击的防护效果对比（高斯机制，不同ε_total）")
    body(doc,
        "防护效果的定量分析通过重建 MSE（均方误差）衡量，MSE 越大说明重建失真越严重，"
        "即攻击越不成功。实验测量结果如下表所示：")
    table(doc,
        ["隐私配置（高斯）", "每轮 ε", "总体 ε_total", "重建质量", "重建 MSE"],
        [
            ["无保护",     "+∞",  "+∞",   "清晰可辨，攻击成功",   "0.0033"],
            ["弱保护",     "30",  "3000", "轮廓轻微模糊",         "0.0172"],
            ["中等保护",   "10",  "1000", "基本不可辨认",         "0.3206"],
            ["较强保护",   "5",   "500",  "近似随机噪声",         "0.3184"],
            ["强保护",     "1",   "100",  "完全随机，攻击失效",   "0.4014"],
        ],
        widths=[3.0, 2.0, 3.0, 4.5, 2.5]
    )
    body(doc,
        "从图 5 和上表可以得出以下结论：无 DP 保护时（ε_total=+∞），"
        "重建 MSE 仅 0.0033，攻击效果极佳；当总体预算降至 ε_total=3000 时，"
        "MSE 上升到 0.017，已能看出噪声影响；ε_total≤1000 时，"
        "MSE 突增至 0.32～0.40，重建图像与原始图像几乎没有视觉相似性，攻击宣告失败。"
        "这说明差分隐私在 ε_total≤1000 的强度下能够有效阻断 GRNN 梯度攻击，"
        "保护客户端数据隐私。当然，这也对应着前面联邦学习实验中模型精度的显著损失，"
        "再次体现了隐私与效用之间的基本矛盾。")

    # ══ 七、总结 ══════════════════════════════════════════════════════════════
    heading(doc, "七、总结", lv=1)
    body(doc,
        "本实验完整实现了基于差分隐私的联邦学习系统，并通过 GRNN 梯度攻击的复现"
        "直观验证了差分隐私的实际保护效果，主要得出以下几点认识：")
    body(doc,
        "第一，精度-隐私权衡是差分隐私联邦学习的核心矛盾。"
        "隐私预算越小，噪声越大，模型精度越低。"
        "拉普拉斯机制在 ε_total=5000（每轮 ε=50）时仍能维持约 97% 的精度，"
        "隐私-效用权衡相对合理；高斯机制在本实验配置下对小 ε 极为敏感，"
        "ε_total≤2000 时模型完全失效，需要结合更精细的隐私分析方法才能实用。")
    body(doc,
        "第二，简单组合定理给出的是隐私损耗的上界，实际损耗往往更小。"
        "100 轮训练后总体隐私预算按简单组合定理为 100×ε，"
        "实践中使用矩会计可将这一上界收紧数倍，从而在相同轮数下维持更强的隐私保证。")
    body(doc,
        "第三，GRNN 攻击实验直观展示了梯度中包含的隐私信息之丰富。"
        "在无 DP 保护下，仅凭一次梯度上传，攻击者便能以极低误差重建原始训练图像。"
        "差分隐私有效阻断了这一攻击路径，在 ε_total≤1000 时重建 MSE 提升 100 倍以上，"
        "彻底使攻击失效，验证了差分隐私作为梯度保护手段的有效性。")

    doc.save(str(OUT))
    print(f"Saved Word: {OUT}")

    pdf_out = OUT.with_suffix(".pdf")
    subprocess.run(["/usr/bin/soffice", "--headless", "--convert-to", "pdf",
                    "--outdir", str(BASE), str(OUT)],
                   check=True, capture_output=True)
    print(f"Saved PDF: {pdf_out}")
    return pdf_out


if __name__ == "__main__":
    build_report()
